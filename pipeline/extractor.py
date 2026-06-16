"""
pipeline/extractor.py

Шаг 1 - Распаковка и инвентаризация (поддерживаемые архивы: ZIP, 7Z, RAR)
Шаг 2 - Извлечение текста:
         PDF     - pymupdf4llm (Markdown)
         DOCX    - python-docx (параграфы + таблицы)
         HTML    - BeautifulSoup
         TXT     - open().read()
         XLSX    - openpyxl
         XLS     - xlrd
         DOC     - antiword (системная утилита)
         RTF     - striprtf

Архивы распаковываются рекурсивно (ZIP внутри 7Z и т.п.), глубина до 3.
"""

import os
import time
import base64
import zipfile
import tempfile
import shutil
import subprocess
import logging
import traceback
from pathlib import Path
from dataclasses import dataclass

import httpx

logger = logging.getLogger(__name__)

# Картинки и .gif/.pptx/.xlsb не парсятся текстом. Картинки идут в Vision-OCR (ниже),
# остальное — пропускаем.
IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".tif", ".tiff", ".bmp", ".webp"}
SKIP_EXTENSIONS = {".pptx", ".xlsb", ".gif"}

MAX_ARCHIVE_DEPTH = 3

# -- Vision-OCR (сканы и картинки через мультимодальную модель АгентПлатформ) --
OCR_ENABLED      = os.getenv("OCR_ENABLED", "1") == "1"
OCR_MODEL        = os.getenv("OCR_MODEL", "google/gemini-2.5-flash")
OCR_MAX_PAGES    = int(os.getenv("OCR_MAX_PAGES", "30"))
OCR_MAX_SIDE_PX  = int(os.getenv("OCR_MAX_SIDE_PX", "2000"))
OCR_API_KEY      = os.getenv("API_KEY", "")
OCR_API_URL      = os.getenv("API_URL", "https://api.agentplatform.ru/v1/chat/completions")
# PDF считаем сканом, если обычное извлечение дало меньше этого числа непробельных символов
PDF_SCAN_MIN_CHARS = int(os.getenv("PDF_SCAN_MIN_CHARS", "200"))
OCR_PROMPT = (
    "Извлеки ВЕСЬ текст из изображения документа дословно, сохраняя структуру "
    "(таблицы передавай текстом по строкам). Только извлечённый текст, без комментариев."
)

# Анти-DoS лимиты на распаковку (см. security-аудит)
ARCHIVE_MAX_TOTAL_BYTES = 500 * 1024 * 1024  # суммарный распакованный размер
ARCHIVE_MAX_FILES = 5000                       # число файлов
ARCHIVE_MAX_RATIO = 120                         # макс. коэффициент сжатия (zip-бомба)


@dataclass
class ExtractedFile:
    name: str
    path: str
    text: str
    doc_type: str = "НЕИЗВЕСТНО"
    skipped: bool = False
    skip_reason: str = ""
    size_bytes: int = 0
    ocr: bool = False        # текст получен через Vision-OCR (скан/картинка)


# -- Архивы (ZIP / 7Z / RAR) ----------------------------------------------

def _fix_zip_filename(name: str) -> str:
    """Исправляет кодировку имён файлов в ZIP (CP437 -> CP866 для кириллицы)."""
    try:
        if any(ord(c) > 127 for c in name):
            fixed = name.encode("cp437").decode("cp866")
            # Есть ли в результате символы кириллицы (диапазон U+0400 .. U+04FF)?
            if any("Ѐ" <= c <= "ӿ" for c in fixed):
                return fixed
    except (UnicodeDecodeError, UnicodeEncodeError):
        pass
    return name


def _unpack_zip_inner(arch_path: Path, target_dir: Path):
    """Чистая распаковка ZIP (без рекурсии). С фиксом кириллических имён,
    защитой от zip-slip (path traversal) и анти-zip-бомба лимитами."""
    base = target_dir.resolve()
    total = 0
    count = 0
    with zipfile.ZipFile(arch_path, "r") as zf:
        for info in zf.infolist():
            # анти zip-бомба: оценка по метаданным без чтения данных
            if info.compress_size and info.file_size / max(info.compress_size, 1) > ARCHIVE_MAX_RATIO:
                logger.warning(f"[ZIP] подозрение на zip-бомбу (ratio), пропуск: {info.filename}")
                continue
            count += 1
            total += info.file_size
            if count > ARCHIVE_MAX_FILES or total > ARCHIVE_MAX_TOTAL_BYTES:
                logger.warning(f"[ZIP] превышен лимит распаковки ({arch_path.name}) — остановка")
                break
            fixed_name = _fix_zip_filename(info.filename)
            dest = (base / fixed_name).resolve()
            # анти zip-slip: dest обязан оставаться внутри base
            if dest != base and base not in dest.parents:
                logger.warning(f"[ZIP] path traversal заблокирован: {info.filename}")
                continue
            if info.is_dir():
                dest.mkdir(parents=True, exist_ok=True)
            else:
                dest.parent.mkdir(parents=True, exist_ok=True)
                dest.write_bytes(zf.read(info.filename))


def _unpack_7z_inner(arch_path: Path, target_dir: Path):
    """Распаковка 7Z через py7zr (чистый Python, без системных утилит)."""
    import py7zr
    with py7zr.SevenZipFile(str(arch_path), mode="r") as z:
        z.extractall(path=str(target_dir))


def _unpack_rar_inner(arch_path: Path, target_dir: Path):
    """Распаковка RAR через unar (свободная утилита, ставится apt-ом)."""
    import rarfile
    rarfile.UNRAR_TOOL = "unar"
    with rarfile.RarFile(str(arch_path)) as rf:
        rf.extractall(path=str(target_dir))


# Маппинг расширения -> функция-распаковщик. Новые форматы — одной строкой.
ARCHIVE_HANDLERS = {
    ".zip": _unpack_zip_inner,
    ".7z":  _unpack_7z_inner,
    ".rar": _unpack_rar_inner,
}


def _unpack_archive(arch_path: Path, target_dir: Path, depth: int = 0) -> list[Path]:
    """Распаковывает архив (ZIP/7Z/RAR) в target_dir. Рекурсивно обрабатывает
    вложенные архивы до глубины MAX_ARCHIVE_DEPTH. Возвращает плоский список
    путей к нераспаковываемым файлам (то есть к фактическим документам).
    """
    if depth > MAX_ARCHIVE_DEPTH:
        logger.warning(f"[ARCH] Превышена глубина вложенности: {arch_path.name}")
        return []

    ext = arch_path.suffix.lower()
    handler = ARCHIVE_HANDLERS.get(ext)
    if not handler:
        return []

    label = ext.upper().lstrip(".")
    try:
        handler(arch_path, target_dir)
        logger.info(f"[{label}] Распакован: {arch_path.name}")
    except Exception as e:
        logger.warning(f"[{label}] Не удалось распаковать {arch_path.name}: {e}")
        return []

    found: list[Path] = []
    for item in target_dir.rglob("*"):
        if not item.is_file():
            continue
        sub_ext = item.suffix.lower()
        if sub_ext in ARCHIVE_HANDLERS:
            nested_dir = item.parent / (item.stem + "_extracted")
            nested_dir.mkdir(exist_ok=True)
            found.extend(_unpack_archive(item, nested_dir, depth + 1))
        else:
            found.append(item)
    return found


# -- Обработчики по форматам ----------------------------------------------

def _extract_pdf(file_path: Path) -> str:
    """PDF -> Markdown через pymupdf4llm."""
    import pymupdf4llm
    logger.info(f"[PDF] pymupdf4llm -> {file_path.name}")
    text = pymupdf4llm.to_markdown(str(file_path))
    logger.info(f"[PDF] {len(text)} симв. из {file_path.name}")
    return text


def _extract_docx(file_path: Path) -> str:
    """DOCX -> текст через python-docx (параграфы + таблицы)."""
    from docx import Document
    logger.info(f"[DOCX] python-docx -> {file_path.name}")
    doc = Document(str(file_path))
    parts: list[str] = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    result = "\n".join(parts)
    logger.info(f"[DOCX] {len(result)} симв. из {file_path.name}")
    return result


def _extract_html(file_path: Path) -> str:
    """HTML -> текст через BeautifulSoup."""
    from bs4 import BeautifulSoup
    logger.info(f"[HTML] BeautifulSoup -> {file_path.name}")
    raw = file_path.read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(raw, "html.parser")
    text = soup.get_text(separator="\n", strip=True)
    logger.info(f"[HTML] {len(text)} симв. из {file_path.name}")
    return text


def _extract_txt(file_path: Path) -> str:
    """TXT -> просто читаем файл."""
    logger.info(f"[TXT] open() -> {file_path.name}")
    text = file_path.read_text(encoding="utf-8", errors="replace")
    logger.info(f"[TXT] {len(text)} симв. из {file_path.name}")
    return text


def _extract_rtf(file_path: Path) -> str:
    """RTF -> текст через striprtf."""
    from striprtf.striprtf import rtf_to_text
    logger.info(f"[RTF] striprtf -> {file_path.name}")
    # RTF обычно в cp1251 или utf-8, striprtf сам разбирается
    raw = file_path.read_text(encoding="utf-8", errors="replace")
    text = rtf_to_text(raw, errors="ignore")
    logger.info(f"[RTF] {len(text)} симв. из {file_path.name}")
    return text


def _extract_xlsx(file_path: Path) -> str:
    """XLSX -> текст через openpyxl."""
    import openpyxl
    logger.info(f"[XLSX] openpyxl -> {file_path.name}")
    wb = openpyxl.load_workbook(str(file_path), read_only=True, data_only=True)
    parts: list[str] = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"--- Лист: {sheet_name} ---")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            line = " | ".join(cells).strip()
            if line.replace("|", "").strip():
                parts.append(line)
    wb.close()
    result = "\n".join(parts)
    logger.info(f"[XLSX] {len(result)} симв. из {file_path.name}")
    return result


def _extract_xls(file_path: Path) -> str:
    """XLS (старый формат) -> текст через xlrd."""
    import xlrd
    logger.info(f"[XLS] xlrd -> {file_path.name}")
    wb = xlrd.open_workbook(str(file_path))
    parts: list[str] = []
    for sheet in wb.sheets():
        parts.append(f"--- Лист: {sheet.name} ---")
        for row_idx in range(sheet.nrows):
            cells = [str(sheet.cell_value(row_idx, col)) for col in range(sheet.ncols)]
            line = " | ".join(cells).strip()
            if line.replace("|", "").strip():
                parts.append(line)
    result = "\n".join(parts)
    logger.info(f"[XLS] {len(result)} симв. из {file_path.name}")
    return result


def _extract_doc_antiword(file_path: Path) -> str:
    """DOC (OLE2) -> текст через antiword. Фолбэк: если .doc — на самом деле HTML."""
    logger.info(f"[DOC] antiword -> {file_path.name}")
    try:
        result = subprocess.run(
            ["antiword", str(file_path)],
            capture_output=True, text=True, timeout=30,
            encoding="utf-8", errors="replace",
        )
        if result.returncode != 0:
            raise RuntimeError(result.stderr.strip() or "antiword вернул ненулевой код")
        text = result.stdout.strip()
        if not text:
            raise RuntimeError("antiword вернул пустой текст")
        logger.info(f"[DOC] {len(text)} симв. из {file_path.name}")
        return text
    except (RuntimeError, FileNotFoundError) as e:
        logger.warning(f"[DOC] antiword не смог: {e}. Пробую как HTML...")
        # Фолбэк: ЕИС часто отдаёт HTML-файлы с расширением .doc
        return _extract_html(file_path)


# Маппинг расширения -> обработчик
FORMAT_HANDLERS = {
    ".pdf":  _extract_pdf,
    ".docx": _extract_docx,
    ".html": _extract_html,
    ".htm":  _extract_html,
    ".txt":  _extract_txt,
    ".rtf":  _extract_rtf,
    ".xlsx": _extract_xlsx,
    ".xls":  _extract_xls,
    ".doc":  _extract_doc_antiword,
}


# -- Vision-OCR -----------------------------------------------------------

def _render_to_jpegs(file_path: Path, is_pdf: bool) -> tuple[list[bytes], bool]:
    """Рендерит документ в JPEG-страницы через PyMuPDF (fitz).
    PDF → каждая страница (апскейл до OCR_MAX_SIDE_PX для читаемости);
    картинка → 1 кадр (только даунскейл крупных). Возвращает (jpegs, обрезано_ли)."""
    import fitz  # PyMuPDF — уже в зависимостях (через pymupdf4llm)
    jpegs: list[bytes] = []
    doc = fitz.open(str(file_path))
    try:
        total = doc.page_count if is_pdf else 1
        n = min(total, OCR_MAX_PAGES)
        for i in range(n):
            page = doc[i]
            long_side = max(page.rect.width, page.rect.height) or 1.0
            if is_pdf:
                scale = OCR_MAX_SIDE_PX / long_side       # рендер PDF в нужный «DPI»
            else:
                scale = min(1.0, OCR_MAX_SIDE_PX / long_side)  # картинку только уменьшаем
            scale = max(0.1, min(scale, 4.0))
            pix = page.get_pixmap(matrix=fitz.Matrix(scale, scale), alpha=False)
            jpegs.append(pix.tobytes("jpeg"))
        truncated = total > OCR_MAX_PAGES
    finally:
        doc.close()
    return jpegs, truncated


def _vision_ocr_image(img_bytes: bytes) -> str:
    """Один vision-запрос: картинка → текст. '' при ошибке (не валит пайплайн)."""
    if not OCR_API_KEY:
        logger.warning("[OCR] API_KEY не задан — vision-OCR пропущен")
        return ""
    b64 = base64.b64encode(img_bytes).decode()
    payload = {
        "model": OCR_MODEL,
        "temperature": 0,
        "max_tokens": 8000,
        "messages": [{"role": "user", "content": [
            {"type": "text", "text": OCR_PROMPT},
            {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64}"}},
        ]}],
    }
    headers = {"Authorization": f"Bearer {OCR_API_KEY}", "Content-Type": "application/json"}
    for attempt in range(1, 4):
        try:
            with httpx.Client(timeout=120.0) as cl:
                r = cl.post(OCR_API_URL, headers=headers, json=payload)
            if r.status_code == 200:
                return (r.json()["choices"][0]["message"]["content"] or "").strip()
            logger.warning(f"[OCR] HTTP {r.status_code}: {r.text[:160]}")
            if r.status_code not in (429, 500, 502, 503, 504):
                break
        except Exception as e:
            logger.warning(f"[OCR] попытка {attempt}: {e}")
        time.sleep(1.0)
    return ""


def _ocr_document(file_path: Path, is_pdf: bool) -> str:
    """Прогоняет документ (скан-PDF или картинку) через Vision-OCR. Склеивает страницы."""
    try:
        jpegs, truncated = _render_to_jpegs(file_path, is_pdf)
    except Exception as e:
        logger.warning(f"[OCR] рендер не удался {file_path.name}: {e}")
        return ""
    if not jpegs:
        return ""
    logger.info(f"[OCR] {file_path.name}: {len(jpegs)} стр. → vision ({OCR_MODEL})")
    parts: list[str] = []
    multi = len(jpegs) > 1
    for idx, img in enumerate(jpegs, 1):
        t = _vision_ocr_image(img)
        if t:
            parts.append((f"\n--- страница {idx} ---\n" if multi else "") + t)
    if truncated:
        parts.append(f"\n[OCR обрезан: обработано первых {OCR_MAX_PAGES} страниц из документа]")
    return "\n".join(parts).strip()


# -- Главная функция ------------------------------------------------------

def extract_files(uploaded_paths: list[Path]) -> list[ExtractedFile]:
    """
    Принимает список путей к загруженным файлам (могут быть архивы).
    Возвращает список ExtractedFile с текстом или пометкой skipped.
    """
    results: list[ExtractedFile] = []
    all_files: list[Path] = []
    tmp_dirs: list[Path] = []

    # --- Шаг 1: инвентаризация ---
    for path in uploaded_paths:
        ext = path.suffix.lower()
        if ext in ARCHIVE_HANDLERS:
            tmp_dir = Path(tempfile.mkdtemp())
            tmp_dirs.append(tmp_dir)
            unpacked = _unpack_archive(path, tmp_dir)
            all_files.extend(unpacked)
            label = ext.upper().lstrip(".")
            logger.info(f"[ИНВЕНТАРИЗАЦИЯ] {label} {path.name}: {len(unpacked)} файлов")
        else:
            all_files.append(path)

    logger.info(f"[ИНВЕНТАРИЗАЦИЯ] Всего файлов: {len(all_files)}")

    # --- Шаг 2: извлечение текста ---
    for file_path in all_files:
        ext = file_path.suffix.lower()
        name = file_path.name
        try:
            size_bytes = file_path.stat().st_size
        except OSError:
            size_bytes = 0

        # Служебные файлы из ZIP (__MACOSX, .DS_Store и т.п.) — тихо игнорируем
        if any(part.startswith("__") or part.startswith(".") for part in file_path.parts):
            logger.info(f"[ПРОПУСК] Служебный: {name}")
            continue

        if ext in SKIP_EXTENSIONS:
            logger.info(f"[ПРОПУСК] {ext}: {name}")
            results.append(ExtractedFile(
                name=name, path=str(file_path), text="", size_bytes=size_bytes,
                skipped=True, skip_reason=f"формат {ext} не поддерживается (презентации/анимация)",
            ))
            continue

        is_image = ext in IMAGE_EXTENSIONS
        handler = FORMAT_HANDLERS.get(ext)
        if not handler and not is_image:
            logger.info(f"[ПРОПУСК] Неизвестный формат {ext}: {name}")
            results.append(ExtractedFile(
                name=name, path=str(file_path), text="", size_bytes=size_bytes,
                skipped=True, skip_reason=f"неизвестный формат {ext}",
            ))
            continue

        try:
            # 1) обычное извлечение текста (для картинок — пусто, у них нет хендлера)
            text = handler(file_path) if handler else ""
            ocr_used = False

            # 2) Vision-OCR там, где текста нет:
            #    - картинка (jpeg/png/...) — всегда;
            #    - PDF-скан — если обычное извлечение дало почти пусто.
            if OCR_ENABLED:
                need_ocr = is_image or (ext == ".pdf" and len(text.strip()) < PDF_SCAN_MIN_CHARS)
                if need_ocr:
                    ocr_text = _ocr_document(file_path, is_pdf=(ext == ".pdf"))
                    if ocr_text.strip():
                        text = ocr_text
                        ocr_used = True
                        logger.info(f"[OCR] {name}: получено {len(ocr_text)} симв. через vision")

            if not text.strip():
                # картинка/скан, который не дался даже OCR — помечаем пропущенным
                if is_image:
                    results.append(ExtractedFile(
                        name=name, path=str(file_path), text="", size_bytes=size_bytes,
                        skipped=True, skip_reason="картинка: OCR не дал текста",
                    ))
                    continue

            if len(text.strip()) < 50:
                logger.warning(f"[ВНИМАНИЕ] {name}: текст короткий ({len(text)} симв.)")
            results.append(ExtractedFile(
                name=name, path=str(file_path), text=text, size_bytes=size_bytes, ocr=ocr_used,
            ))
        except Exception as e:
            logger.error(f"[ОШИБКА] {name}: {e}")
            logger.error(traceback.format_exc())
            results.append(ExtractedFile(
                name=name, path=str(file_path), text="", size_bytes=size_bytes,
                skipped=True, skip_reason=f"ошибка чтения: {str(e)[:120]}",
            ))

    # Очистка временных директорий
    for tmp_dir in tmp_dirs:
        shutil.rmtree(tmp_dir, ignore_errors=True)

    ok = sum(1 for r in results if not r.skipped)
    skip = sum(1 for r in results if r.skipped)
    logger.info(f"[ИТОГ] Обработано: {len(results)}, успешно: {ok}, ошибки: {skip}")
    return results
